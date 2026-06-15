from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "assets" / "CV_Xinlong_Le_EN.docx"


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=9.5, bold=bold)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "D9DEE5")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12 if level == 1 else 8, after=4)
    r = p.add_run(text)
    set_run_font(r, size=13 if level == 1 else 11.5, bold=True, color=BLUE if level == 1 else DARK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=3, line=1.12)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    run = p.add_run(text)
    set_run_font(run, size=10.2)
    return p


def add_pub(doc, text, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    set_paragraph_spacing(p, before=0, after=4, line=1.1)
    run = p.add_run(text)
    set_run_font(run, size=9.8)
    if note:
        r2 = p.add_run(" " + note)
        set_run_font(r2, size=9.4, color=MUTED)
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=2, line=1.0)
    r = title.add_run("Xinlong Le (ShiloH)")
    set_run_font(r, size=22, bold=True, color=RGBColor(0, 0, 0))

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact, after=7, line=1.0)
    r = contact.add_run(
        "Ph.D. Candidate in Hydraulic Engineering | Huazhong University of Science and Technology\n"
        "Email: xinlong_le@hust.edu.cn | Phone: +86 183 8229 0228 | ORCID: 0000-0002-8681-1036"
    )
    set_run_font(r, size=9.5, color=MUTED)

    add_heading(doc, "Research Profile")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=5, line=1.12)
    r = p.add_run(
        "Ph.D. candidate with interdisciplinary training in hydrology, climate extremes, remote sensing "
        "and multi-source Earth observation data assessment. Research experience covers watershed "
        "hydrological responses under global change, compound extreme events, precipitation product "
        "evaluation, CMIP6 scenario data, statistical downscaling, dataset construction and uncertainty diagnosis."
    )
    set_run_font(r, size=10.3)

    add_heading(doc, "Education")
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Period", "Institution", "Degree / Major"]):
        set_cell_text(cell, text, bold=True)
    rows = [
        ("2023-2027", "Huazhong University of Science and Technology", "Ph.D., Hydraulic Engineering"),
        ("2020-2023", "Southwest Jiaotong University", "M.S., Tunnel Engineering"),
        ("2016-2020", "Hainan University", "B.E., Road and Bridge Engineering"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)

    add_heading(doc, "Research Interests")
    for item in [
        "Watershed hydrological responses under global change",
        "Compound climate extremes and hydrometeorological risk",
        "Satellite, reanalysis and climate model data evaluation and fusion",
        "Statistical downscaling, bias correction and uncertainty analysis",
        "Climate datasets, extreme indices and Earth system data products",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Selected Publications")
    pubs = [
        ("Le Xinlong, Ling Kang*, et al. Multi-Criteria Evaluation and Ranking of 28 Precipitation Products over China. International Journal of Applied Earth Observation and Geoinformation, 2026.", "[JCR Q1, CAS Q1 TOP]"),
        ("Le Xinlong, Ling Kang*, et al. ECHIDNA: Extreme Climate Historical and Future Indices Data under Numerous Approaches across Major Chinese River Basins Based on CMIP6 Multi-Model Ensemble. Scientific Data, 2026.", "[JCR Q1, CAS Q2 TOP]"),
        ("Le Xinlong, Kang Ling*, et al. Evaluating statistical downscaling and bias-correction methods for climate extremes across major Chinese River basins. Ain Shams Engineering Journal, 2026.", "[JCR Q1]"),
        ("Xinlong Le, Kang Ling, et al. A novel hybrid biological optimisation algorithm for tackling reservoir optimal operation problem. Ain Shams Engineering Journal, 2025.", "[JCR Q1]"),
        ("Le Xinlong, Kang Ling*, et al. Quantifying the impacts of human activities and climate change on water resource changes: a case study of Hubei Province. 41st IAHR World Congress, 2025.", "[Conference paper and oral presentation]"),
        ("Hao Chen, Ling Kang*, Le Xinlong, et al. River System Thermal Dynamics under Dual Pressures of Climate Change and Cascade Reservoir Operations. Journal of Environmental Management, 2025.", "[JCR Q1, CAS Q2 TOP]"),
        ("Ling Kang, Le Xinlong*, et al. Merging Multi-Source Precipitation Products: A Global Review of Datasets and Methods with Insights from the Tibetan Plateau. Earth-Science Reviews.", "[JCR Q1, CAS Q1 TOP]"),
        ("Le Xinlong, Kang Ling*, et al. Multiscale Applicability of PERSIANN-Family Satellite Precipitation Products across Global Land Areas. International Journal of Applied Earth Observation and Geoinformation, 2026.", "[Under review]"),
        ("Le Xinlong, Ling Kang*, et al. A ranking-correction-fusion framework (HARMONY) for climate-zone-aware multi-source daily precipitation over China at 0.1 degrees. Ain Shams Engineering Journal, 2026.", "[Under review]"),
        ("Le Xinlong, Ling Kang*, et al. PREMISE v2.1: A configuration-driven decision-support framework for multi-source precipitation product evaluation and ranking. Computers and Geosciences, 2026.", "[Under review]"),
    ]
    for text, note in pubs:
        add_pub(doc, text, note)

    add_heading(doc, "Patents and Software Copyright")
    for item in [
        "Five invention patents related to water-resource attribution, reservoir operation, compound dry-hot event identification, tunnel deformation simulation and tunnel health monitoring.",
        "Two utility model patents related to tunnel loading experiments and model-test auxiliary devices.",
        "Software copyright: Multi-objective joint optimal operation system for cascade reservoirs.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Research Projects")
    for item in [
        "Fundamental Research Funds for the Central Universities, 2023-2026, Principal Investigator.",
        "China Yangtze Power research project on cascade reservoir digital operation and scheduling, 2023-2026, responsible member.",
        "National Key R&D Program on basin-scale flood scenario deduction and intelligent decision-making, 2022-2025, participant.",
        "National Key R&D Program on integrated water-system decision technologies for the Yellow River Basin, 2023-2027, participant.",
        "Hubei water conservancy key research project on digital-twin reservoir operation, 2024-2025, responsible member.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Academic Activities")
    for item in [
        "Oral presentation at the 41st IAHR World Congress in Singapore.",
        "Invited reviewer for the 5th International Conference on Artificial Intelligence, Information Processing and Cloud Computing (AIIPCC 2025).",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Technical Skills")
    for item in [
        "Python and Java for hydrometeorological data processing, scientific computing, model evaluation and visualization.",
        "NetCDF, GeoTIFF, station observations, remote sensing products, reanalysis datasets and multi-dimensional gridded data.",
        "CMIP6 scenario data, statistical downscaling, bias correction, climate-extreme indices and uncertainty diagnosis.",
        "Hydrological modeling, precipitation product evaluation, multi-source data fusion, multi-criteria decision analysis, machine learning and swarm-intelligence optimization.",
        "Scientific writing, reviewer response, oral presentation and interdisciplinary collaboration.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Awards and Honors")
    for item in [
        "Science and Technology Innovation Scholarship, 2025.",
        "First-Class Doctoral Scholarship, 2023 and 2024.",
        "First-Class Academic Scholarship, 2018, 2020 and 2021; Special Academic Scholarship, 2019.",
        "First Prize, Challenge Cup Reveal-the-List Competition, 2025.",
        "Third Prize, Huawei Cup Graduate Mathematical Contest in Modeling, 2023.",
        "Outstanding Graduate and Merit Student, 2020; Outstanding Communist Youth League Member, 2018.",
    ]:
        add_bullet(doc, item)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer, after=0, line=1.0)
    fr = footer.add_run("Xinlong Le (ShiloH) | Curriculum Vitae | Updated June 2026")
    set_run_font(fr, size=8.5, color=MUTED)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
